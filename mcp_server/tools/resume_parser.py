"""
mcp_server/tools/resume_parser.py
────────────────────────────────────
Extracts raw text from a resume (PDF or DOCX) and detects skills present
in it, organised into categories for nicer display in the skill-gap
learning plan.

Reuses the single skill taxonomy already defined in jd_gap_analyzer.py
so resume-side and JD-side skill detection never drift apart.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from mcp_server.tools.jd_gap_analyzer import _ALL_SKILLS, _extract_skills_from_text

# ── Category buckets (display grouping only — detection itself is unified) ───

_CATEGORIES: dict[str, set[str]] = {
    "languages": {
        "python", "java", "javascript", "typescript", "scala", "go", "rust",
        "c++", "c#", "r", "sql", "bash", "shell", "php", "kotlin", "swift",
    },
    "data_engineering": {
        "pyspark", "spark", "hive", "hadoop", "kafka", "airflow", "dbt",
        "flink", "pandas", "numpy", "polars", "databricks", "snowflake",
    },
    "ml_ai": {
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "mlflow",
        "vertex ai", "sagemaker", "hugging face", "langchain", "openai",
    },
    "cloud": {
        "aws", "gcp", "azure", "cloud run", "lambda", "ec2", "s3", "bigquery",
        "cloud functions", "cloud storage", "gke", "eks",
    },
    "devops": {
        "docker", "kubernetes", "k8s", "terraform", "ansible", "helm",
        "github actions", "jenkins", "ci/cd", "argocd",
    },
    "web": {
        "react", "angular", "vue", "django", "flask", "fastapi", "express",
        "nextjs", "graphql", "rest", "grpc",
    },
    "databases": {
        "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
        "cassandra", "dynamodb", "firestore",
    },
    "tools": {
        "git", "linux", "unix", "jira", "confluence", "postman", "grafana",
        "prometheus", "looker", "tableau", "power bi",
    },
    "enterprise": {"sas", "spss", "stata", "sap", "oracle"},
    "business_soft": {
        "agile", "scrum", "product management", "data analysis", "communication",
    },
}

# Skills in _ALL_SKILLS but not yet bucketed above fall back to "other" —
# this guard keeps the two files honest if the taxonomy grows later.
_BUCKETED = set().union(*_CATEGORIES.values())
_UNBUCKETED = _ALL_SKILLS - _BUCKETED


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_docx(path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_pdf(path)
    if suffix in (".docx", ".dotx"):
        return _extract_text_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported resume format: {suffix} (expected .pdf, .docx, or .txt)")


# ── Skill detection ────────────────────────────────────────────────────────────

def _categorize_skills(found: set[str]) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {cat: sorted(found & skills) for cat, skills in _CATEGORIES.items()}
    other = sorted(found & _UNBUCKETED)
    if other:
        out["other"] = other
    # Drop empty categories so downstream display isn't cluttered.
    return {cat: skills for cat, skills in out.items() if skills}


def _detect_years_experience(text: str) -> int | None:
    """Best-effort extraction of total years of experience, if stated explicitly."""
    matches = re.findall(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+(?:of\s+)?experience", text, re.I)
    if not matches:
        return None
    return int(float(max(matches, key=float)))


# ── Public API ────────────────────────────────────────────────────────────────

def parse_resume(resume_path: str) -> dict[str, Any]:
    """
    Parse a resume file and return its raw text plus categorised detected skills.

    Parameters
    ----------
    resume_path : Path to a .pdf, .docx, or .txt resume file.

    Returns
    -------
    On success, a dict with:
        raw_text          – full extracted plain text
        detected_skills   – dict of category -> sorted list of skills found
        skill_count       – total number of distinct skills detected
        years_experience  – best-effort int, or None if not explicitly stated
        word_count        – word count of raw_text (rough length signal)
    On failure, a dict with a single "error" key — callers check for this
    before proceeding (matches the convention used by jd_gap_analyzer / ats_scorer).
    """
    path = Path(resume_path)

    if not path.exists():
        return {"error": f"Resume file not found: {resume_path}"}

    try:
        raw_text = _extract_text(path)
    except ValueError as e:
        return {"error": str(e)}
    except Exception as e:
        return {"error": f"Failed to read resume: {e}"}

    if not raw_text.strip():
        return {
            "error": (
                "No extractable text found in resume — it may be a scanned/image-only "
                "PDF. Try exporting a text-based version, or OCR it first."
            )
        }

    found_skills = _extract_skills_from_text(raw_text)
    detected_skills = _categorize_skills(found_skills)

    return {
        "raw_text": raw_text,
        "detected_skills": detected_skills,
        "skill_count": len(found_skills),
        "years_experience": _detect_years_experience(raw_text),
        "word_count": len(raw_text.split()),
    }
